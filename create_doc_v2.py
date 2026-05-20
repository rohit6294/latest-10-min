from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SS  = r'C:\Users\rohit\Desktop\10  min\screenshots'
OUT = r'C:\Users\rohit\Desktop\10  min\Suraksha Kavach_Product_Document_v2.docx'

RED   = RGBColor(0xC0, 0x00, 0x00)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
NAVY  = RGBColor(0x0F, 0x3D, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0x55, 0x55, 0x66)
GREEN = RGBColor(0x16, 0x70, 0x2E)
AMBER = RGBColor(0xB4, 0x5A, 0x00)

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shd(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex_color)
    tcPr.append(s)

def para(text='', bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=6, italic=False):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text); r.bold=bold; r.italic=italic
        r.font.size=Pt(size); r.font.color.rgb = color or DARK
    return p

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(8)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(22)
    r.font.color.rgb=RED; r.font.all_caps=True

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=NAVY

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=DARK

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'C00000')
    pBdr.append(bot); pPr.append(pBdr)

def bullet(text, bold_prefix=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(0.7)
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
    r0=p.add_run('•  '); r0.font.color.rgb=RED; r0.font.size=Pt(10)
    if bold_prefix and text.startswith(bold_prefix):
        r1=p.add_run(bold_prefix); r1.bold=True; r1.font.size=Pt(10); r1.font.color.rgb=DARK
        r2=p.add_run(text[len(bold_prefix):]); r2.font.size=Pt(10); r2.font.color.rgb=color or GRAY
    else:
        r1=p.add_run(text); r1.font.size=Pt(10); r1.font.color.rgb=color or DARK

def img(path, caption='', width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p=doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(12)
        r=p.runs[0]; r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GRAY

def callout(text, title='', bg='F0F4FF'):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shd(c,bg); c.width=Cm(15)
    p=c.paragraphs[0]
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.left_indent=Cm(0.3); p.paragraph_format.right_indent=Cm(0.3)
    if title:
        r=p.add_run(f'{title}  '); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=NAVY
    r2=p.add_run(text); r2.font.size=Pt(10); r2.font.color.rgb=DARK
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def table(headers, rows, col_widths=None):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; shd(c,'1A1A2E')
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for ri,row in enumerate(rows):
        dr=t.rows[ri+1]; bg='F9F9F9' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=dr.cells[ci]; shd(c,bg)
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(9); r.font.color.rgb=DARK
    if col_widths:
        for row in t.rows:
            for ci,c in enumerate(row.cells): c.width=col_widths[ci]
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

def step_block(code, num, title, desc):
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    c0=t.cell(0,0); c1=t.cell(0,1); shd(c0,'1A1A2E')
    c0.width=Cm(2.2); c1.width=Cm(13.3)
    p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before=Pt(6); p0.paragraph_format.space_after=Pt(2)
    r0=p0.add_run(code); r0.bold=True; r0.font.size=Pt(8); r0.font.color.rgb=RGBColor(255,128,128)
    p0b=c0.add_paragraph(); p0b.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p0b.paragraph_format.space_after=Pt(6)
    r0b=p0b.add_run(str(num)); r0b.bold=True; r0b.font.size=Pt(24); r0b.font.color.rgb=WHITE
    p1=c1.paragraphs[0]; p1.paragraph_format.left_indent=Cm(0.3)
    p1.paragraph_format.space_before=Pt(6)
    r1=p1.add_run(title); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=NAVY
    p2=c1.add_paragraph(); p2.paragraph_format.left_indent=Cm(0.3)
    p2.paragraph_format.space_after=Pt(8)
    r2=p2.add_run(desc); r2.font.size=Pt(10); r2.font.color.rgb=GRAY
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Suraksha Kavach'); r.bold=True; r.font.size=Pt(44); r.font.color.rgb=RED

p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=p2.add_run('Emergency Ambulance Dispatch Platform'); r2.font.size=Pt(18); r2.font.color.rgb=NAVY

doc.add_paragraph()
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=p3.add_run('Complete Product Overview — For Non-Technical Readers')
r3.font.size=Pt(13); r3.italic=True; r3.font.color.rgb=GRAY

for _ in range(2): doc.add_paragraph()
t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
c=t.cell(0,0); shd(c,'1A1A2E')
p4=c.paragraphs[0]; p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before=Pt(14); p4.paragraph_format.space_after=Pt(14)
r4=p4.add_run('"Get an Ambulance to Any Emergency in Under 10 Minutes"')
r4.bold=True; r4.font.size=Pt(13); r4.font.color.rgb=WHITE
for _ in range(4): doc.add_paragraph()

mt=doc.add_table(rows=4,cols=2); mt.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([('Document Type','Product & Stakeholder Brief'),
    ('Version','v2.0 — May 2026'),('Audience','Non-Technical Stakeholders, Investors, Partners'),
    ('Status','Final — For Distribution')]):
    c0,c1=mt.rows[i].cells; shd(c0,'F0F0F0'); shd(c1,'FFFFFF')
    r0=c0.paragraphs[0].add_run(k); r0.bold=True; r0.font.size=Pt(9); r0.font.color.rgb=NAVY
    r1=c1.paragraphs[0].add_run(v); r1.font.size=Pt(9); r1.font.color.rgb=DARK
    c0.width=Cm(5); c1.width=Cm(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  01 EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1('01. Executive Summary'); divider()
para('Suraksha Kavach is a technology-powered emergency ambulance coordination platform '
     'that connects people facing medical emergencies to the nearest available '
     'ambulance driver and hospital — in under 10 minutes. It works across three '
     'connected tools: a public website anyone can open on any phone, an Android '
     'app for ambulance drivers, and a web dashboard for hospitals.',size=11,sa=8)
para('This document explains everything about the platform — what it is, why it '
     'was built, how every screen works, how the live map tracks the ambulance, '
     'how hospitals see and manage their bed availability, how the driver app '
     'guides a driver from alert to hospital — written for anyone, with no '
     'technical knowledge assumed.',size=11,sa=8)
callout('Suraksha Kavach is live and operational. Every screen shown in this document '
        'is from the actual, running product.',title='KEY STATEMENT:',bg='FFF0F0')

# ══════════════════════════════════════════════════════════════════════════════
#  02 PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1('02. Problem Statement'); divider()
para('Every year, thousands of people die not because medicine could not save '
     'them, but because an ambulance arrived too late. In India, the emergency '
     'response system faces deep, structural problems:',size=11,sa=8)
h2('The Five Core Problems')
for t,d in [
    ('No single dispatch system.','  Families call 5–10 hospitals one by one during a crisis. '
     'There is no central system that finds the nearest ambulance automatically.'),
    ('Drivers have no navigation.','  Ambulance drivers depend on a panicked family '
     'member directing them over the phone, causing delays and wrong turns.'),
    ('Hospitals have no advance notice.','  A hospital only knows an ambulance is '
     'coming when it arrives at the gate. No team is ready, no bed is prepared.'),
    ('Zero visibility for families.','  Once help is called, the family has no way '
     'to track where the ambulance is or how long it will take.'),
    ('WhatsApp and calls are uncoordinated.','  Many people try to reach help via '
     'WhatsApp or phone, but there is no structured system to capture and route '
     'these requests to a real driver.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(4)
    r1=p.add_run('• '); r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(t); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(d); r3.font.size=Pt(10); r3.font.color.rgb=GRAY
para()
h2('Before vs. After Suraksha Kavach')
table(['Situation','Without Suraksha Kavach','With Suraksha Kavach'],[
    ['Cardiac arrest at home','Family calls 5+ numbers, waits 30–60 min','Driver notified in seconds, arrives <10 min'],
    ['Hospital bed availability','Unknown until ambulance arrives at gate','Checked live before dispatch'],
    ['Driver navigation','Guided by panicked family on phone','Automatic GPS route to patient'],
    ['Family during crisis','Helpless, no visibility','Live tracking link on phone'],
    ['Hospital preparation','No advance notice, team not ready','Pre-alerted with patient details'],
],col_widths=[Cm(4.5),Cm(5.5),Cm(5.5)])

# ══════════════════════════════════════════════════════════════════════════════
#  03 PURPOSE
# ══════════════════════════════════════════════════════════════════════════════
h1('03. Purpose'); divider()
callout('Ensure that no one dies waiting for an ambulance because the system '
        'failed to connect them to help in time.',title='OUR PURPOSE:',bg='FFF0F0')
para('Suraksha Kavach fills the gap between the moment an emergency happens and '
     'the moment professional help arrives. Every feature in the app — from '
     'the SOS button to the live map to the hospital bed panel — exists to '
     'make that gap as small as possible.',size=11,sa=8)

# ══════════════════════════════════════════════════════════════════════════════
#  04 OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
h1('04. Objectives'); divider()
for t,d in [
    ('Reduce ambulance response time to under 10 minutes',' in urban areas.'),
    ('Give families real-time tracking',' from the moment help is dispatched.'),
    ('Pre-alert hospitals automatically',' so beds and teams are ready before arrival.'),
    ('Create a verified driver network',' matched to the right ambulance type for each emergency.'),
    ('Make emergency access multi-channel',' — web, WhatsApp, or phone call.'),
    ('Give hospitals a live bed management dashboard',' updated in real time.'),
    ('Build an admin control layer',' to monitor every active emergency and intervene if needed.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(4)
    r1=p.add_run('• '); r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(t); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(d); r3.font.size=Pt(10); r3.font.color.rgb=GRAY

# ══════════════════════════════════════════════════════════════════════════════
#  05 WHAT IS IT & WHY
# ══════════════════════════════════════════════════════════════════════════════
h1('05. What Is Suraksha Kavach?'); divider()
para('Suraksha Kavach is three connected tools that work as one system:',size=11,sa=6)
for title,desc in [
    ('The Public Website (10minrescue.com)',
     'Open on any phone or computer — no app download needed. Anyone can press '
     '"Emergency SOS", describe what happened, share their location, and within '
     'seconds the system finds the nearest driver. Also accessible via WhatsApp '
     'or a phone callback request.'),
    ('The Driver Android App',
     'Ambulance drivers carry this on their phones. When a nearby emergency '
     'matches their ambulance type, they receive an immediate alert with a '
     '30-second countdown. They see the patient location, emergency details, '
     'and pre-selected hospital, then navigate directly to the patient.'),
    ('The Hospital Web Dashboard',
     'Hospitals log in on any computer or tablet. They see every ambulance '
     'assigned to them in real time, manage their bed availability with + / − '
     'buttons, and view the full history of completed trips.'),
]:
    h3(title); para(desc,size=10,sa=8,color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
#  06 AMBULANCE TYPES & SEATING
# ══════════════════════════════════════════════════════════════════════════════
h1('06. Ambulance Types — Equipment & Seating Layout'); divider()
para('Suraksha Kavach matches each emergency to the right type of ambulance. '
     'There are three ambulance types, each designed for a different level '
     'of medical need. When a patient selects their urgency level (Critical, '
     'Serious, or Stable), the system automatically recommends the correct '
     'ambulance type — though the patient can override this if needed.',
     size=11,sa=10)

# Type A
t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
c=t.cell(0,0); shd(c,'FFF0F0'); c.width=Cm(15.5)
p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
p.paragraph_format.left_indent=Cm(0.3)
r=p.add_run('TYPE A  —  ICU AMBULANCE   (For: CRITICAL emergencies)')
r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RED
doc.add_paragraph().paragraph_format.space_after=Pt(0)

para('Used for:  Cardiac arrest, unconscious patient, no pulse or breathing, '
     'life-threatening trauma.',size=10,sb=2,sa=4,color=DARK)
h3('Equipment on board:')
for item in ['Cardiac monitor and defibrillator (restarts stopped heart)',
             'Ventilator (breathes for the patient if they cannot breathe)',
             'Pulse oximeter (monitors blood oxygen levels)',
             'IV drip lines and emergency medications',
             'Oxygen cylinder with mask and flow meter',
             'Suction machine (clears airways)',
             'Spinal board and neck collar (for trauma cases)',
             'AED — Automated External Defibrillator']:
    bullet(item)
h3('Seating & space layout:')
para('The ICU ambulance has the most equipment-dense interior:',size=10,sa=4,color=GRAY)
for item in ['1 stretcher (patient lies flat, fully monitored)',
             '1 paramedic seat alongside the stretcher (for constant monitoring)',
             '1 attendant fold-down seat at the patient\'s feet',
             'All equipment mounted on walls and ceiling rails for quick access',
             'No standing room — every inch is used for life-support equipment']:
    bullet(item)
callout('In a cardiac arrest, the paramedic in a Type A ambulance can administer '
        'defibrillation shocks and run an IV drip while the ambulance is moving — '
        'this is life-saving care that cannot wait until hospital arrival.',
        title='WHY IT MATTERS:',bg='FFF0F0')

para()
# Type B
t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
c=t.cell(0,0); shd(c,'FFFDE7'); c.width=Cm(15.5)
p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
p.paragraph_format.left_indent=Cm(0.3)
r=p.add_run('TYPE B  —  ADVANCED AMBULANCE   (For: SERIOUS emergencies)')
r.bold=True; r.font.size=Pt(13); r.font.color.rgb=AMBER
doc.add_paragraph().paragraph_format.space_after=Pt(0)

para('Used for:  Severe injuries, heavy bleeding, breathing difficulties, '
     'high-risk pregnancies, major road accidents.',size=10,sb=2,sa=4,color=DARK)
h3('Equipment on board:')
for item in ['Oxygen cylinder and mask',
             'Defibrillator (for heart rhythm correction)',
             'Advanced monitoring — blood pressure, ECG, oxygen saturation',
             'IV drip lines and fluids',
             'First aid and wound management supplies',
             'Immobilisation equipment — splints, collars, backboard',
             'Emergency medications — pain relief, anti-seizure drugs']:
    bullet(item)
h3('Seating & space layout:')
para('The Advanced ambulance has more flexibility than the ICU type:',size=10,sa=4,color=GRAY)
for item in ['1 stretcher (patient lies flat or is semi-reclined)',
             '1 EMT (Emergency Medical Technician) seat alongside the patient',
             '1–2 attendant seats for a family member or second responder',
             'Storage compartments for medication and equipment along both walls',
             'More accessible interior — EMT can move around the patient']:
    bullet(item)

para()
# Type C
t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
c=t.cell(0,0); shd(c,'F1F8E9'); c.width=Cm(15.5)
p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
p.paragraph_format.left_indent=Cm(0.3)
r=p.add_run('TYPE C  —  NORMAL AMBULANCE   (For: STABLE / transport cases)')
r.bold=True; r.font.size=Pt(13); r.font.color.rgb=GREEN
doc.add_paragraph().paragraph_format.space_after=Pt(0)

para('Used for:  Minor injuries requiring hospital transport, post-discharge '
     'transfers, stable patients who need monitoring en route.',size=10,sb=2,sa=4,color=DARK)
h3('Equipment on board:')
for item in ['Basic oxygen supply',
             'First aid kit (bandages, wound care)',
             'Stretcher or wheelchair depending on patient mobility',
             'Basic vitals monitoring (blood pressure, pulse)']:
    bullet(item)
h3('Seating & space layout:')
for item in ['1 stretcher or reclining seat for the patient',
             '2 attendant seats — typically for a family member and a first-aider',
             'Spacious interior — designed for comfort over intensive care',
             'Fold-down step at the rear for easy boarding by mobile patients']:
    bullet(item)

para()
h2('How the System Selects the Ambulance Type Automatically')
table(['Urgency','Ambulance Type','Typical Condition','Who Rides Along'],[
    ['CRITICAL (Red)','Type A — ICU','Cardiac arrest, no pulse, unconscious','Paramedic with life support skills'],
    ['SERIOUS (Amber)','Type B — Advanced','Severe injury, heavy bleeding, breathing issues','EMT with advanced first aid training'],
    ['STABLE (Green)','Type C — Normal','Minor injury, transport needed','First-aider or driver only'],
],col_widths=[Cm(3.2),Cm(3.8),Cm(5.0),Cm(3.5)])
callout('The patient or their family can override the auto-selected type on the '
        'SOS wizard (Step 2). If they are unsure, "Let System Decide" is always '
        'the safest choice.',title='NOTE:',bg='F0F4FF')

# ══════════════════════════════════════════════════════════════════════════════
#  07 HOW IT SOLVES THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
h1('07. How Suraksha Kavach Solves the Problem'); divider()
h2('The Complete Emergency Journey — Step by Step')
para('Here is exactly what happens from the moment someone calls for help to '
     'the moment the patient is safe in hospital:',size=11,sa=10)
steps=[
    ('STEP 1',1,'Someone Sends an SOS',
     'The patient\'s family opens 10minrescue.com on any phone. They tap '
     '"Emergency SOS" and fill in three things: what happened (e.g. "father '
     'collapsed, chest pain"), how serious it is (Critical/Serious/Stable), '
     'and the system auto-selects the ambulance type. They share their GPS '
     'location and hit "SEND SOS NOW". The whole process takes under 60 seconds.'),
    ('STEP 2',2,'System Finds the Nearest Driver',
     'Instantly, the system searches for available drivers within 1 km of the '
     'patient. Only drivers with the matching ambulance type are notified. '
     'If no driver responds in 30 seconds, the search expands to 2 km, then '
     '5 km, then 10 km — automatically, with no human intervention needed.'),
    ('STEP 3',3,'Driver Gets an Alert on Their Phone',
     'The driver\'s Android app shows a full-screen emergency alert with a '
     '30-second countdown timer. They see the patient\'s name, phone number, '
     'emergency description, the urgency level (colour-coded), the ambulance '
     'type required, and the pre-selected hospital. They tap ACCEPT to respond.'),
    ('STEP 4',4,'Driver Navigates to the Patient',
     'After accepting, the driver app shows the patient\'s exact GPS location '
     'on a map. The family immediately receives a live tracking link — they '
     'can watch the ambulance icon moving toward them in real time.'),
    ('STEP 5',5,'Hospital Is Pre-Alerted',
     'The moment a hospital is confirmed (either by the patient during SOS or '
     'by the driver after pickup), the hospital\'s dashboard shows the incoming '
     'ambulance. Hospital staff see the patient\'s condition, ambulance type, '
     'and estimated arrival time. They prepare the right bed and the right team.'),
    ('STEP 6',6,'Patient Is Picked Up & Transported',
     'The driver marks "Patient Picked Up" in the app. The hospital sees this '
     'status update immediately. The ambulance navigates to the hospital while '
     'the patient receives care en route. The family can still track the vehicle.'),
    ('STEP 7',7,'Patient Arrives, Hospital Is Ready',
     'The driver marks "Arrived at Hospital". The hospital logs the patient '
     'intake. The bed is ready, the team is at the door. The trip is marked '
     'complete and the entire journey is stored in the system.'),
]
for s in steps: step_block(s[0],s[1],s[2],s[3])

# ══════════════════════════════════════════════════════════════════════════════
#  08 UI SCREENS — WEBSITE
# ══════════════════════════════════════════════════════════════════════════════
h1('08. Application Screens — Complete Visual Walkthrough'); divider()
para('Every screen in the Suraksha Kavach platform is shown below with a plain-language '
     'explanation of what it does and who uses it.',size=11,sa=10)

h2('8.1  The Public Website — What Patients & Families See')
h3('Home Page / Landing Page')
para('The first page anyone sees at 10minrescue.com. Two large buttons dominate '
     'the screen: "Emergency SOS" (red, for active emergencies) and '
     '"Request Callback" (for when someone needs us to call them). The page '
     'also explains how the service works and builds trust.',size=10,sa=6,color=GRAY)
img(os.path.join(SS,'01_landing_hero_desktop.png'),
    'Figure 1: Suraksha Kavach Home Page — Emergency SOS and Request Callback are immediately visible',
    width=Inches(5.8))

h3('"How It Works" Section')
para('Scrolling down the home page shows the three ways to reach Suraksha Kavach: '
     'the SOS button on the website, a callback request, or WhatsApp/phone call. '
     'Every channel connects to the same dispatch system.',size=10,sa=6,color=GRAY)
img(os.path.join(SS,'06_how_it_works.png'),
    'Figure 2: Three channels — Emergency SOS, Callback Request, and WhatsApp/Call',
    width=Inches(5.8))

h3('"From SOS to Hospital" Flow')
para('This section on the website shows the patient the complete journey in '
     'four steps: Patient sends SOS → Driver accepts → Hospital is pre-alerted '
     '→ Live tracking begins. This builds confidence that the system handles '
     'everything automatically.',size=10,sa=6,color=GRAY)
img(os.path.join(SS,'07_features.png'),
    'Figure 3: End-to-end real-time coordination — Patient → Driver → Hospital → Live Track',
    width=Inches(5.8))

# SOS Wizard
h2('8.2  The SOS Wizard — Step-by-Step Emergency Request')
para('When someone taps Emergency SOS, they enter a guided 4-step wizard. '
     'A progress bar at the top shows exactly where they are in the process. '
     'The design is intentionally simple — large buttons, minimal text, '
     'colour-coded choices — so it works even under extreme stress.',
     size=10,sa=8,color=GRAY)

h3('Step 1 — Describe the Emergency & Select Urgency')
para('The user types a brief description of what happened (e.g. "Father collapsed, '
     'chest pain, not breathing"). They then select one of three urgency levels:',
     size=10,sa=4,color=DARK)
for label,col in [
    ('CRITICAL (Red)','cardiac arrest, unconscious, not breathing, life at immediate risk'),
    ('SERIOUS (Amber)','severe injury, heavy bleeding, breathing difficulties, major trauma'),
    ('STABLE (Green)','minor injury, patient is conscious, needs transport to hospital'),
]:
    bullet(f'{label}: {col}')
img(os.path.join(SS,'02_sos_step1_mobile.png'),
    'Figure 4: SOS Step 1 — Emergency description and three urgency levels (Critical/Serious/Stable)',
    width=Inches(2.8))

h3('Step 2 — Ambulance Type (Auto-Selected, Overridable)')
para('Based on the urgency level, the system automatically highlights the right '
     'ambulance type: Type A (ICU) for Critical, Type B (Advanced) for Serious, '
     'Type C (Normal) for Stable. The user can change this if needed.',
     size=10,sa=6,color=GRAY)

h3('Step 3 — Choose a Hospital (or Let the System Decide)')
para('The system shows nearby hospitals sorted by rating and distance, filtered '
     'to only those with available beds of the required type. Each hospital card '
     'shows the hospital name, distance, rating, and number of available beds. '
     'If the patient selects "Let System Decide", the driver will choose the '
     'best hospital after picking up the patient.',size=10,sa=6,color=GRAY)

h3('Step 4 — Review All Details and Send SOS')
para('A final summary screen shows every detail: the emergency description, '
     'urgency level, ambulance type, hospital choice, and GPS coordinates. '
     'The large red "SEND SOS NOW" button submits everything in one tap. '
     'This is the final confirmation — nothing is sent by accident.',
     size=10,sa=6,color=GRAY)

h3('Callback Request Page')
para('For users who cannot type (elderly, very distressed, or low internet), '
     'the callback page lets them enter just their name and phone number. '
     'An operator calls them back immediately to coordinate help.',
     size=10,sa=6,color=GRAY)
img(os.path.join(SS,'08_callback.png'),
    'Figure 5: Callback Request — drop name and phone number, we call you back',
    width=Inches(2.8))

# ══════════════════════════════════════════════════════════════════════════════
#  09 DRIVER APP SCREENS
# ══════════════════════════════════════════════════════════════════════════════
h1('09. Driver App — Every Screen Explained'); divider()
para('The Suraksha Kavach Driver App runs on Android smartphones. It is designed '
     'for use in a moving vehicle — large buttons, colour-coded urgency, minimal '
     'reading required. Below is every screen a driver interacts with during '
     'a normal shift.',size=11,sa=10)

h2('9.1  Driver Home Screen — Online/Offline Status & Live Alerts')
para('This is the first screen a driver sees after logging in. It has three '
     'main areas:',size=10,sa=6,color=DARK)
for t,d in [
    ('SOS Alert Card (top of screen, red background):',
     '  When a nearby emergency matches the driver\'s ambulance type, a red '
     'alert card appears instantly at the top. It shows the emergency description, '
     'urgency level, distance to patient, ambulance type required, and pre-selected '
     'hospital. A large "ACCEPT REQUEST" button is prominently displayed.'),
    ('Status Hero (middle):',
     '  Shows the driver\'s current status — Online & Available (green dot '
     'pulsing) or Offline. A toggle switch lets them go online or offline. '
     'When online, their GPS location is shared with the system every few seconds. '
     'The ambulance type they are assigned to (e.g. "Type A — ICU Ambulance") '
     'is shown as a chip below the status.'),
    ('Quick Stats Grid (bottom):',
     '  Four stat cards show: Today\'s Trips completed, distance to nearest '
     'pending request, total trips completed all time, and the driver\'s '
     'average customer rating out of 5.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(5)
    r1=p.add_run('• '); r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(t); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(d); r3.font.size=Pt(10); r3.font.color.rgb=GRAY
img(os.path.join(SS,'D1_driver_home.png'),
    'Figure 6: Driver Home Screen — Online status, live SOS alert, stats grid',
    width=Inches(2.8))
callout('The driver also has a side menu (hamburger icon) with: My Profile, '
        'Trip History, Document Upload, and Sign Out.',title='SIDE MENU:',bg='F0F4FF')

h2('9.2  Incoming Request Alert — The 30-Second Decision Screen')
para('When a new emergency comes in and the driver has been selected, the entire '
     'phone screen changes to a full-screen alert. This screen is designed to '
     'force immediate attention:',size=10,sa=6,color=DARK)
for t,d in [
    ('Countdown Timer (top centre):',
     '  A circular countdown shows how many seconds the driver has to respond. '
     'It starts at 30 seconds. If the driver does not respond, the request '
     'automatically goes to the next nearest driver.'),
    ('Urgency Gradient Background:',
     '  The entire screen changes colour based on urgency — deep red gradient '
     'for Critical, amber for Serious, green for Stable. This is impossible '
     'to miss even in a moving vehicle.'),
    ('Pulsing Emergency Icon:',
     '  A large ambulance/siren icon pulses and scales in and out to demand '
     'attention. The urgency label (CRITICAL / SERIOUS / STABLE) appears '
     'below it in a white badge, with the ambulance type below that.'),
    ('Patient Information Card:',
     '  A frosted card shows: Patient name, phone number, emergency description '
     '(up to 2 lines), and the pre-selected hospital name.'),
    ('Two Buttons at the Bottom:',
     '  "Decline" (outlined, white) and "ACCEPT REQUEST" (solid white with '
     'red/amber/green text). If the driver accepts, they immediately go to '
     'the navigation screen. If another driver has already accepted the same '
     'request, the app shows a snackbar: "Request already taken."'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(5)
    r1=p.add_run('• '); r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(t); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(d); r3.font.size=Pt(10); r3.font.color.rgb=GRAY
img(os.path.join(SS,'D2_driver_incoming.png'),
    'Figure 7: Incoming Request Alert — 30-second countdown, CRITICAL urgency, patient info, Accept/Decline',
    width=Inches(2.8))

h2('9.3  Navigation to Patient — Getting to the Emergency')
para('After accepting, the driver immediately sees the navigation screen. '
     'It shows a live map (OpenStreetMap) with:',size=10,sa=6,color=DARK)
for item in [
    'A blue marker showing the driver\'s own position (updates in real time)',
    'A red pulsing marker showing the patient\'s exact GPS location',
    'A blue route line drawn along actual roads between the driver and patient',
    'Turn-by-turn navigation guidance',
    'Patient name, phone number, and emergency description shown below the map',
    'A "Call Patient" button for direct phone contact if needed',
    'A "Mark Patient Picked Up" button — tapped when the driver arrives and loads the patient',
]:
    bullet(item)
callout('The route line is calculated in real time using road data. As the driver '
        'moves, both the route and the ETA update automatically.',title='HOW THE MAP WORKS:',bg='F0F4FF')

h2('9.4  Navigation to Hospital')
para('After marking "Patient Picked Up", the screen switches to hospital navigation:',
     size=10,sa=6,color=DARK)
for item in [
    'Live map now shows the route from current location to the pre-selected hospital',
    'Hospital name, address, and phone number shown below the map',
    'A "Call Hospital" button for advance coordination',
    'Patient\'s condition and ambulance type shown as a reminder',
    '"Confirm Arrival at Hospital" button — tapped when the ambulance reaches the hospital gate',
]:
    bullet(item)
para('When the driver confirms arrival, the hospital dashboard immediately shows '
     'the status as "Arrived". The trip is then marked complete.',
     size=10,sa=6,color=GRAY)

h2('9.5  Driver Registration & Document Upload')
para('New drivers who want to join the Suraksha Kavach network register through the '
     'app. The registration screen collects:',size=10,sa=6,color=DARK)
for item in ['Full name and phone number',
             'Ambulance type (A, B, or C) they operate',
             'Vehicle registration number',
             'Driving licence number',
             'Document photo uploads: licence, vehicle permit, insurance certificate',
             'Selfie photo for identity verification']:
    bullet(item)
para('Once submitted, the admin team reviews the documents and either approves '
     'or rejects the driver. Until approved, the driver cannot go online or '
     'receive requests. If rejected, they see the reason and can resubmit.',
     size=10,sa=6,color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
#  10 LIVE TRACKING MAP
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Live Tracking Page — How the Map Works'); divider()
para('After the SOS is sent, the patient or family immediately gets a live '
     'tracking link. Opening it shows a real-time map of the ambulance. No '
     'app download is needed — it works in any mobile browser.',
     size=11,sa=10)

h2('What the Patient\'s Family Sees on the Tracking Page')
img(os.path.join(SS,'T1_live_tracking.png'),
    'Figure 8: Live Tracking Page — ambulance position, ETA, driver info, hospital, emergency call button',
    width=Inches(2.8))

h2('Every Element on the Tracking Page — Explained')
for t,d in [
    ('Status Banner (top):',
     '  Shows the current status in large text — "Finding nearest driver..." '
     '(amber) while searching, then "Driver is on the way" (green) once '
     'accepted. The colour and icon change automatically as the status changes.'),
    ('Distance and ETA boxes:',
     '  Once a driver is assigned, two boxes appear showing the current '
     'distance from the driver to the patient (e.g. 1.4 km) and the '
     'estimated arrival time (e.g. ~3 min). Both update live as the driver moves.'),
    ('The Live Map:',
     '  A real map (OpenStreetMap) showing two markers and a route line:'),
    ('  🚑  Blue ambulance marker','  = the driver\'s position, updating every few seconds'),
    ('  📍  Red pulsing marker','  = the patient\'s exact GPS location (where they sent SOS from)'),
    ('  Blue route line','  = the actual road route between driver and patient'),
    ('Driver Information Card:',
     '  Shows the driver\'s name, vehicle number, and ambulance type. '
     'A green phone button lets the family call the driver directly.'),
    ('Hospital Information Card:',
     '  Shows the pre-selected hospital\'s name and address. '
     'A blue phone button lets the family call the hospital directly.'),
    ('Emergency Call Strip (bottom):',
     '  A red button always visible at the bottom — "+91 78660 67136 — '
     'Emergency Line" — so the family can call us directly at any point.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(5)
    r1=p.add_run('• '); r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(t); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(d); r3.font.size=Pt(10); r3.font.color.rgb=GRAY

h2('How the Map Updates in Real Time')
para('The live map is powered by a real-time database. Here is exactly what '
     'happens behind the scenes (explained without technical jargon):',
     size=10,sa=6,color=DARK)
for i,step in enumerate([
    'Every few seconds, the driver\'s phone sends its GPS coordinates to our server.',
    'The tracking page "listens" for any change — like a subscription to live news.',
    'The moment new coordinates arrive, the ambulance icon on the map moves to the new position.',
    'The route line is redrawn along actual roads using real road-routing data.',
    'The distance and ETA numbers recalculate based on the new position.',
    'This all happens automatically — the family does not need to refresh the page.',
],1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
    r1=p.add_run(f'{i}.  '); r1.bold=True; r1.font.color.rgb=RED; r1.font.size=Pt(10)
    r2=p.add_run(step); r2.font.size=Pt(10); r2.font.color.rgb=DARK

# ══════════════════════════════════════════════════════════════════════════════
#  11 HOSPITAL DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Hospital Dashboard — Every Section Explained'); divider()
para('Every partner hospital gets a secure login to the Hospital Portal. '
     'It is a web-based dashboard that works on any computer or tablet. '
     'No app download required.',size=11,sa=10)

h2('11.1  Hospital Login Screen')
para('Hospitals log in with an email and password issued by the Suraksha Kavach '
     'admin team. The login page is separate from the patient website, '
     'accessible only at /hospital.',size=10,sa=6,color=GRAY)
img(os.path.join(SS,'03_hospital_login_desktop.png'),
    'Figure 9: Hospital Portal Login — secure, role-specific access',
    width=Inches(4.5))

h2('11.2  The Full Hospital Dashboard — After Login')
img(os.path.join(SS,'H1_hospital_dashboard.png'),
    'Figure 10: Hospital Dashboard — Stats, Bed Availability Panel, Tabs, and Request Cards',
    width=Inches(5.8))

h2('11.3  Statistics Bar (Top of Dashboard)')
para('Four live counters at the top update automatically:',size=10,sa=6,color=DARK)
table(['Counter','What It Shows','Colour Signal'],[
    ['Incoming','Number of ambulances currently assigned to this hospital and en route',
     'Red background + pulsing dot when count > 0'],
    ['In Transit','Ambulances currently carrying a patient to this hospital',
     'Blue — steady'],
    ['Completed Today','Trips successfully completed at this hospital today',
     'Green — steady'],
    ['Total Requests','All-time requests assigned to this hospital','Navy — steady'],
],col_widths=[Cm(3.2),Cm(8.0),Cm(4.3)])

h2('11.4  Bed Availability Panel — Real-Time Bed Management')
para('This is the most important section of the hospital dashboard. It shows '
     'exactly how many beds of each type are available right now, and lets '
     'staff update the count with a single tap.',size=10,sa=8,color=DARK)
para('There are three bed type rows, each colour-coded:',size=10,sa=4,color=DARK)
for label,col,type_,who in [
    ('ICU Beds (red)','For patients requiring intensive care and life-support equipment',
     'Type A','Critical emergencies — cardiac, respiratory failure'),
    ('Advanced Beds (amber)','For patients requiring monitoring and advanced first aid',
     'Type B','Serious emergencies — trauma, accidents, high-risk maternity'),
    ('Normal Beds (green)','For stable patients requiring basic hospital admission',
     'Type C','Minor injuries, post-procedure transfers'),
]:
    h3(label)
    bullet(f'Purpose: {col}')
    bullet(f'Linked to: {type_} ambulances')
    bullet(f'Used for: {who}')

h3('How to update bed count:')
para('Each row shows the current available count and the total bed capacity '
     '(e.g. "4 / 7 total"). A progress bar shows the ratio visually. '
     'Staff use two buttons:',size=10,sa=4,color=GRAY)
for t,d in [
    ('  −  (Minus button — Red):','  Tap when a patient is admitted and occupies a bed. '
     'This reduces the available count by 1 and immediately updates what patients '
     'see when choosing a hospital during an SOS.'),
    ('  +  (Plus button — Green):','  Tap when a patient is discharged and the '
     'bed is available again. This increases the count by 1.'),
]:
    bullet(f'{t}{d}')
callout('Patients see this bed count in real time when choosing a hospital during '
        'the SOS wizard. A hospital with zero beds of the required type will NOT '
        'appear as a choice — ensuring ambulances are only sent where beds exist.',
        title='WHY THIS MATTERS:',bg='FFF0F0')
para('Hospitals can also tap "Set Total Beds" to configure their total capacity '
     '(e.g. "We have 10 ICU beds total"). This only needs to be done once '
     'or when the hospital adds or removes beds.',size=10,sa=8,color=GRAY)

h2('11.5  The Three Tabs — Incoming, In Transit, History')
para('Below the bed panel, three tabs organise all ambulance requests:',
     size=10,sa=6,color=DARK)
table(['Tab','What It Shows','When to Check'],[
    ['Incoming','Ambulances confirmed for this hospital, currently en route',
     'Constantly — these need immediate preparation'],
    ['In Transit','Ambulances with the patient on board, travelling to hospital',
     'When preparing for imminent arrivals'],
    ['History','All completed and cancelled trips — searchable archive',
     'For reporting, auditing, and performance review'],
],col_widths=[Cm(3.0),Cm(7.5),Cm(5.0)])

h2('11.6  Request Cards — What Each Card Shows')
para('Each ambulance request appears as a card. An incoming request card '
     '(highlighted in red) shows:',size=10,sa=4,color=DARK)
for item in ['Emergency type icon (e.g. ❤️ for cardiac, 🚗 for accident, 🤱 for maternity)',
             'Patient name and phone number',
             'Patient\'s GPS location coordinates',
             'Emergency category (Cardiac, Accident, Maternity, General, etc.)',
             'Time since the request was created (e.g. "2m ago")',
             '"View Details" button — opens a full detail view with all request information']:
    bullet(item)
callout('Hospitals do NOT need to accept or reject ambulances. They are automatically '
        'notified. Their job is to prepare the right team and bed before the '
        'ambulance arrives.',title='HOW IT WORKS:',bg='F0F4FF')

# ══════════════════════════════════════════════════════════════════════════════
#  12 ADMIN PANEL
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Admin Panel — Operations Control Centre'); divider()
para('The Admin Panel is restricted to Suraksha Kavach\'s own operations team. '
     'It gives full visibility and control over the entire platform.',
     size=11,sa=8)
img(os.path.join(SS,'05_admin_login.png'),
    'Figure 11: Admin Panel Login — authorised personnel only',
    width=Inches(4.5))
para('Admin capabilities include:',size=10,sa=4,color=DARK)
for item in ['Review and approve/reject driver registrations and uploaded documents',
             'Monitor all active rescue requests across the city in real time',
             'See which drivers are currently online and where they are',
             'Manage hospital accounts and configure bed totals',
             'View full trip history and performance reports']:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
#  13 EDGE CASES
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Edge Cases — What Happens When Things Go Wrong?'); divider()
para('A life-saving platform must handle every failure. Here is how '
     'Suraksha Kavach responds to each one:',size=11,sa=10)
edge_cases=[
    ('No driver available nearby',
     'Search automatically expands every 30 seconds: 1 km → 2 km → 5 km → 10 km. '
     'If still no driver, the admin team is alerted and can manually coordinate.'),
    ('Driver does not accept within 30 seconds',
     'Request automatically moves to the next nearest available driver. '
     'This repeats until someone accepts.'),
    ('Patient\'s GPS location cannot be detected',
     'The SOS form shows a "tap Allow for location" prompt. If still unavailable, '
     'the patient can use the Callback page or WhatsApp to describe their address verbally.'),
    ('Hospital bed count reaches zero',
     'That hospital disappears from the patient\'s hospital selection list. '
     'Only hospitals with confirmed available beds appear as options.'),
    ('Driver loses internet mid-journey',
     'The last known location is cached. Navigation continues offline. '
     'The app re-syncs as soon as the connection returns.'),
    ('Patient closes browser after sending SOS',
     'The SOS is already submitted and cannot be lost. The SOS ID is saved '
     'in the browser so the patient can return to the tracking page.'),
    ('Multiple SOS requests from the same location',
     'The system detects duplicates from the same GPS coordinates within '
     'a short time window and prevents multiple drivers being dispatched.'),
    ('Hospital does not update bed count',
     'Admins see hospitals with stale counts. Bed counts older than 24 hours '
     'are flagged as "unverified" to avoid incorrect dispatch.'),
    ('Another driver accepts the same request first',
     'The second driver\'s app shows: "Request already taken by another driver" '
     'and returns them to the home screen automatically.'),
    ('WhatsApp message is incomplete',
     'The WhatsApp system asks follow-up questions to collect location and '
     'urgency before routing to a driver.'),
]
for title,desc in edge_cases:
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    c0=t.cell(0,0); c1=t.cell(0,1); shd(c0,'FFF0F0')
    c0.width=Cm(0.5); c1.width=Cm(15.0)
    p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before=Pt(4)
    r0=p0.add_run('!'); r0.bold=True; r0.font.size=Pt(11); r0.font.color.rgb=RED
    p1=c1.paragraphs[0]; p1.paragraph_format.left_indent=Cm(0.3)
    p1.paragraph_format.space_before=Pt(4); p1.paragraph_format.space_after=Pt(2)
    r1a=p1.add_run(title+':  '); r1a.bold=True; r1a.font.size=Pt(10); r1a.font.color.rgb=DARK
    r1b=p1.add_run(desc); r1b.font.size=Pt(10); r1b.font.color.rgb=GRAY
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
#  14 SCOPE / OUT OF SCOPE
# ══════════════════════════════════════════════════════════════════════════════
h1('14. Scope — What Suraksha Kavach Covers'); divider()
for item in [
    'Emergency SOS via web browser (mobile and desktop)',
    'Automatic nearby driver search with GPS-based geolocation',
    'Smart radius expansion (1 km → 10 km) when no driver is immediately found',
    'Three ambulance types matched to three urgency levels',
    'Live hospital bed availability check during the SOS flow',
    'Live ambulance tracking for patients and families (real-time map)',
    'Driver Android app: online/offline toggle, incoming alerts, GPS navigation',
    'Hospital web dashboard: live request view, bed management, trip history',
    'Admin web panel: driver verification, platform monitoring, hospital management',
    'WhatsApp emergency intake via Cloud API',
    'Callback request form for voice-based coordination',
    'Push notification system for drivers (FCM)',
    'Real-time database — all request statuses, driver locations, bed counts',
    'Role-based access: patient, driver, hospital, admin',
    'Full trip history and status lifecycle recording',
]:
    bullet(item,color=DARK)

h1('15. Out of Scope'); divider()
for title,desc in [
    ('Payment processing','No billing or invoicing is integrated. Pricing is handled offline.'),
    ('iOS Driver App','The driver app supports Android only. iOS is not part of this release.'),
    ('In-app chat','Communication is by phone call. No in-app messaging feature.'),
    ('Government ambulance integration','The platform operates its own verified network.'),
    ('Automatic fare estimation','Trip costs are handled separately, not in-app.'),
    ('Medical records','Patient history and prescriptions are not stored.'),
    ('Fire or police dispatch','The platform is ambulance-specific only.'),
    ('AI medical advice','The system collects urgency info but never provides diagnoses.'),
    ('Multi-city rollout (v1)','Current version is deployed for a single city.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(4)
    r1=p.add_run('✗  '); r1.bold=True; r1.font.size=Pt(10); r1.font.color.rgb=RED
    r2=p.add_run(title+': '); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=DARK
    r3=p.add_run(desc); r3.font.size=Pt(10); r3.font.color.rgb=GRAY

# ══════════════════════════════════════════════════════════════════════════════
#  16 SUCCESS MATRIX
# ══════════════════════════════════════════════════════════════════════════════
h1('16. Success Matrix'); divider()
para('Success is measured across three areas: speed, reliability, and reach.',
     size=11,sa=10)
h2('Response & Speed Metrics')
table(['Metric','What It Measures','Target'],[
    ['Average Response Time','Time from SOS sent to driver accepting','< 3 minutes'],
    ['Ambulance On-Scene Time','Time from SOS sent to ambulance at patient','< 10 minutes'],
    ['Driver Accept Rate','% of requests accepted by first notified driver','> 80%'],
    ['Hospital Pre-Alert Rate','% of trips where hospital is notified before arrival','100%'],
    ['SOS Completion Rate','% of SOS requests resulting in a completed trip','> 90%'],
],col_widths=[Cm(4.8),Cm(7.0),Cm(3.7)])
h2('User Experience Metrics')
table(['Metric','Target'],[
    ['SOS form completion time (Steps 1–4)','< 60 seconds'],
    ['Family satisfaction score (post-trip feedback)','> 4.2 / 5.0'],
    ['Hospital dashboard active adoption','> 85% of partner hospitals'],
    ['Driver daily active usage','> 70% of registered drivers go online daily'],
    ['Callback response time','< 5 minutes'],
],col_widths=[Cm(9.0),Cm(6.5)])
h2('Network Growth (6-Month Targets)')
table(['Metric','Target'],[
    ['Active verified drivers','50+ drivers online weekly'],
    ['Partner hospitals with active dashboards','20+ hospitals'],
    ['Total SOS requests served','500+ cases'],
    ['City coverage','Full urban area covered'],
    ['Average driver rating','> 4.0 / 5.0'],
],col_widths=[Cm(9.0),Cm(6.5)])

# ══════════════════════════════════════════════════════════════════════════════
#  17 CLOSING
# ══════════════════════════════════════════════════════════════════════════════
h1('17. Final Note'); divider()
para('Suraksha Kavach is built on one belief: in a medical emergency, every minute '
     'without help is a minute too long. The platform exists to make those '
     'minutes disappear.',size=11,sa=8)
para('This document has covered everything — what the platform is, why it was '
     'built, how every screen works, how the map tracks the ambulance in real '
     'time, how hospitals manage their bed availability live, and how the '
     'driver\'s phone guides them from alert to patient to hospital.',
     size=11,sa=8)
callout('Email: mcteam.myewards@gmail.com  |  Phone: +91 78660 67136  |  '
        'Website: 10minrescue.com',title='CONTACT:',bg='F0F4FF')
doc.add_paragraph()
para('Suraksha Kavach — Product Brief v2.0 — May 2026',size=8,color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
